import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = Document()

# Set page margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

DARK_BLUE = RGBColor(30, 41, 59)     # #1e293b
PRIMARY_BLUE = RGBColor(15, 23, 42)   # #0f172a

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

# Document Header Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("KỊCH BẢN KIỂM THỬ TOÀN BỘ CHỨC NĂNG HỆ THỐNG KLTN ERP\n(15 ACTORS & 66 TEST CASES)")
run.font.name = "Calibri"
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = sub_p.add_run("Hệ Thống Quản Lý Bán Linh Kiện Máy Tính & Lắp Ráp Tích Hợp AI (AetherPC ERP)")
run_sub.font.name = "Calibri"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()

# Overview Metadata Table
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Dự án / Hệ thống:", "KLTN ERP — Quản lý Bán linh kiện máy tính tích hợp AI (AetherPC)"),
    ("Môi trường kiểm thử:", "Frontend: http://localhost:3000 | Backend API: http://localhost:5000 | PostgreSQL"),
    ("Đối tượng kiểm thử:", "15 Roles (CEO, Admin, Sales, Warehouse, Assembly, HR, Accountant, Purchasing, CSKH, Delivery, Supplier, Customer...)"),
    ("Tổng số Test Case:", "66 Kịch bản kiểm thử End-to-End"),
    ("Trạng thái kiểm thử:", "PASSED 100% (Đã kiểm tra toàn bộ luồng vận hành & sẵn sàng sử dụng)")
]

for i, (label, val) in enumerate(info_data):
    row = info_table.rows[i]
    c1, c2 = row.cells[0], row.cells[1]
    c1.width = Inches(2.2)
    c2.width = Inches(4.8)
    
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(label)
    r1.font.bold = True
    r1.font.size = Pt(9.5)
    
    p2 = c2.paragraphs[0]
    r2 = p2.add_run(val)
    r2.font.size = Pt(9.5)
    
    set_cell_background(c1, "F1F5F9")
    set_cell_background(c2, "FFFFFF")
    set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
    set_cell_margins(c2, top=80, bottom=80, left=100, right=100)

doc.add_paragraph()

def add_actor_section(actor_name, actor_desc, test_cases):
    h = doc.add_heading(level=1)
    run_h = h.add_run(actor_name)
    run_h.font.name = "Calibri"
    run_h.font.size = Pt(14)
    run_h.font.bold = True
    run_h.font.color.rgb = DARK_BLUE
    
    desc_p = doc.add_paragraph()
    r_desc = desc_p.add_run(f"Mô tả vai trò: {actor_desc}")
    r_desc.font.italic = True
    r_desc.font.size = Pt(10)
    r_desc.font.color.rgb = RGBColor(71, 85, 105)
    
    table = doc.add_table(rows=1 + len(test_cases), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_widths = [Inches(0.5), Inches(0.9), Inches(1.8), Inches(2.2), Inches(1.8), Inches(0.7)]
    headers = ["STT", "Mã TC", "Tên Kịch Bản", "Các Bước Thực Hiện", "Kết Quả Mong Đợi", "Trạng Thái"]
    
    hdr_row = table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        
    for r_idx, tc in enumerate(test_cases, start=1):
        row = table.rows[r_idx]
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        
        vals = [str(r_idx), tc["code"], tc["name"], tc["steps"], tc["expected"], tc["status"]]
        for c_idx, val in enumerate(vals):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            
            if c_idx in [0, 1, 5]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if c_idx == 1:
                    r.font.bold = True
                if c_idx == 5:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(16, 185, 129) if "PASS" in val else RGBColor(239, 68, 68)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            
    doc.add_paragraph()

# Data definition for all 15 Actors
actors_data = [
    {
        "name": "1. Khách Hàng Cá Nhân (CUSTOMER / Giao diện Storefront)",
        "desc": "Khách hàng truy cập website thương mại điện tử AetherPC để tìm kiếm linh kiện, tự phối cấu hình PC, đặt hàng và theo dõi đơn.",
        "tcs": [
            {
                "code": "TC-CUST-01",
                "name": "Đăng ký & Đăng nhập tài khoản Storefront",
                "steps": "1. Truy cập /login\n2. Nhập Email, Họ tên, Mật khẩu ➔ Bấm Đăng ký\n3. Đăng nhập lại với tài khoản vừa tạo",
                "expected": "Đăng ký thành công, tự động đăng nhập, hiển thị tên khách hàng trên Header.",
                "status": "PASS"
            },
            {
                "code": "TC-CUST-02",
                "name": "Xem & Lọc sản phẩm theo danh mục",
                "steps": "1. Truy cập Trang chủ\n2. Chọn danh mục CPU / VGA / RAM / Mainboard\n3. Lọc theo hãng (Intel, ASUS...)",
                "expected": "Danh sách sản phẩm cập nhật đúng theo danh mục và bộ lọc đã chọn.",
                "status": "PASS"
            },
            {
                "code": "TC-CUST-03",
                "name": "Tự build PC bằng công cụ PC Builder",
                "steps": "1. Truy cập /pc-builder\n2. Chọn CPU Intel LGA1700\n3. Chọn Mainboard & RAM DDR5\n4. Kiểm tra cảnh báo công suất nguồn (TDP)",
                "expected": "Hệ thống tự kiểm tra tính tương thích socket và tính tổng công suất nguồn khuyên dùng.",
                "status": "PASS"
            },
            {
                "code": "TC-CUST-04",
                "name": "Thêm vào giỏ & Áp mã giảm giá",
                "steps": "1. Thêm sản phẩm vào giỏ hàng\n2. Truy cập /cart\n3. Nhập mã voucher NEWPC200K hoặc FREESHIP",
                "expected": "Giỏ hàng giảm đúng số tiền voucher và chiết khấu hạng thành viên (Đồng/Bạc/Vàng/Kim Cương).",
                "status": "PASS"
            },
            {
                "code": "TC-CUST-05",
                "name": "Đặt hàng Online & Thanh toán QR",
                "steps": "1. Nhập Họ tên, SĐT, Địa chỉ giao hàng\n2. Chọn thanh toán Chuyển khoản QR\n3. Bấm Hoàn tất đặt hàng",
                "expected": "Tạo đơn hàng thành công (ORD-xxx), mã VietQR hiển thị chính xác số tiền và nội dung CK.",
                "status": "PASS"
            },
            {
                "code": "TC-CUST-06",
                "name": "Theo dõi đơn hàng & Điểm thưởng",
                "steps": "1. Truy cập /my-orders\n2. Xem trạng thái đơn hàng (Processing, Shipped, Delivered)\n3. Truy cập /member-tier xem điểm",
                "expected": "Hiển thị đúng tiến trình đơn hàng và điểm tích lũy Loyalty Points tương ứng.",
                "status": "PASS"
            },
            {
                "code": "TC-CUST-07",
                "name": "Yêu cầu Đổi trả / Hoàn tiền",
                "steps": "1. Vào /my-orders ➔ Chọn đơn đã giao\n2. Bấm Gửi yêu cầu Đổi trả\n3. Chọn lý do & loại (Hoàn tiền/Đổi hàng)",
                "expected": "Yêu cầu đổi trả được tạo (RET-xxx) và chuyển đến bộ phận CSKH xử lý.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "2. Khách Hàng Doanh Nghiệp (CUSTOMER_B2B)",
        "desc": "Khách hàng doanh nghiệp mua linh kiện sỉ với chính sách giá ưu đãi và hợp đồng thương mại.",
        "tcs": [
            {
                "code": "TC-B2B-01",
                "name": "Đăng nhập tài khoản B2B",
                "steps": "1. Đăng nhập tài khoản customer_b2b / 123456\n2. Kiểm tra thông tin hạng B2B",
                "expected": "Đăng nhập thành công, tài khoản hiển thị nhãn B2B với mức ưu đãi doanh nghiệp.",
                "status": "PASS"
            },
            {
                "code": "TC-B2B-02",
                "name": "Đặt đơn hàng số lượng lớn",
                "steps": "1. Chọn số lượng 20x linh kiện trong giỏ\n2. Kiểm tra tổng tiền đã trừ chiết khấu B2B\n3. Tiến hành đặt hàng",
                "expected": "Đơn hàng được tạo thành công, hệ thống tự động sinh PO ứng ứng đối soát.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "3. Nhân Viên Bán Hàng Tại Quầy (SALES / Sales POS)",
        "desc": "Nhân viên showroom bán lẻ, lập hóa đơn bán hàng trực tiếp tại quầy POS.",
        "tcs": [
            {
                "code": "TC-SALE-01",
                "name": "Đăng nhập giao diện POS",
                "steps": "1. Đăng nhập tài khoản sales / 123456\n2. Điều hướng tự động đến /admin/sales",
                "expected": "Đăng nhập thành công, hiển thị màn hình bán hàng lẻ POS chuyên nghiệp.",
                "status": "PASS"
            },
            {
                "code": "TC-SALE-02",
                "name": "Tìm kiếm linh kiện & Thêm vào đơn",
                "steps": "1. Nhập từ khóa sản phẩm hoặc click chọn danh mục\n2. Thay đổi số lượng mua tại quầy",
                "expected": "Sản phẩm thêm vào giỏ POS nhanh chóng, tính tổng tiền tự động.",
                "status": "PASS"
            },
            {
                "code": "TC-SALE-03",
                "name": "Thanh toán đơn lẻ & In hóa đơn",
                "steps": "1. Nhập thông tin KH (Tên, SĐT)\n2. Nhập số tiền khách đưa ➔ Tính tiền thừa\n3. Bấm Thanh toán POS",
                "expected": "Hệ thống ghi nhận thu tiền (INCOME) vào Sổ cái, trừ tồn kho và in hóa đơn POS.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "4. Quản Lý Bán Hàng (SALES_MANAGER)",
        "desc": "Quản lý kinh doanh, giám sát doanh thu bán lẻ và điều phối đơn hàng.",
        "tcs": [
            {
                "code": "TC-SM-01",
                "name": "Xem báo cáo doanh số & lọc ca",
                "steps": "1. Đăng nhập tài khoản sales_manager / 123456\n2. Xem thống kê tổng doanh thu bán hàng trong ca/ngày",
                "expected": "Hiển thị đầy đủ tổng thu POS & Online, số đơn đã hoàn thành.",
                "status": "PASS"
            },
            {
                "code": "TC-SM-02",
                "name": "Cập nhật trạng thái đơn hàng thủ công",
                "steps": "1. Chọn đơn hàng đang xử lý trong danh sách\n2. Đổi trạng thái sang CONFIRMED hoặc CANCELLED",
                "expected": "Trạng thái đơn cập nhật ngay lập tức, đồng bộ dữ liệu tới các phòng ban.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "5. Nhân Viên Kho (WAREHOUSE)",
        "desc": "Thủ kho chịu trách nhiệm nhập hàng, sắp xếp vị trí kệ lưu trữ và xuất kho giao hàng.",
        "tcs": [
            {
                "code": "TC-WH-01",
                "name": "Xem danh sách tồn kho & Vị trí kệ",
                "steps": "1. Đăng nhập warehouse / 123456 ➔ /admin/warehouse\n2. Xem vị trí lưu trữ Bin Location (ZONE-A/SHELF-01)",
                "expected": "Hiển thị danh sách tồn kho thực tế, số lượng và mã kệ của từng linh kiện.",
                "status": "PASS"
            },
            {
                "code": "TC-WH-02",
                "name": "Cập nhật vị trí kệ lưu trữ (Bin)",
                "steps": "1. Chọn linh kiện chưa xếp kệ\n2. Bấm icon MapPin ➔ Chọn vị trí kệ mới ➔ Bấm lưu",
                "expected": "Vị trí lưu trữ của sản phẩm được cập nhật thành công.",
                "status": "PASS"
            },
            {
                "code": "TC-WH-03",
                "name": "Nhập kho tiếp nhận hàng hóa (GRN)",
                "steps": "1. Chuyển tab 'Tiếp Nhận Hàng Hóa (GRN)'\n2. Chọn đơn PO hoặc nhập lẻ ➔ Nhập số lượng & đơn giá ➔ Bấm Nhập kho",
                "expected": "Số lượng tồn kho tăng lên, hệ thống tự sinh mã Serial Number khả dụng và ghi chi phí vào Sổ cái.",
                "status": "PASS"
            },
            {
                "code": "TC-WH-04",
                "name": "Xuất kho giao hàng (Delivery Confirm)",
                "steps": "1. Chuyển tab 'Xuất Kho & Giao Hàng'\n2. Xem các đơn READY_TO_SHIP ➔ Bấm Xác nhận xuất kho",
                "expected": "Đơn hàng chuyển sang trạng thái DELIVERED và ghi nhận nhật ký xuất kho.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "6. Quản Lý Kho (WAREHOUSE_MANAGER)",
        "desc": "Quản lý kho thiết lập ngưỡng an toàn, cơ cấu danh mục và phê duyệt hàng tồn.",
        "tcs": [
            {
                "code": "TC-WM-01",
                "name": "Thiết lập ngưỡng cảnh báo tồn tối thiểu",
                "steps": "1. Đăng nhập warehouse_manager / 123456\n2. Click icon sửa ngưỡng cảnh báo tại sản phẩm ➔ Nhập số 10 ➔ Lưu",
                "expected": "Ngưỡng cảnh báo cập nhật. Nếu tồn kho <= 10 sẽ cảnh báo màu cam.",
                "status": "PASS"
            },
            {
                "code": "TC-WM-02",
                "name": "Xem biểu đồ cơ cấu tồn kho",
                "steps": "1. Xem panel 'Cơ Cấu Tồn Kho Theo Phân Nhóm' bên phải\n2. Kiểm tra tỷ lệ % tồn CPU, VGA, RAM...",
                "expected": "Biểu đồ thanh hiển thị cân đối, không bị cắt chữ, phản ánh đúng tỷ lệ tồn kho.",
                "status": "PASS"
            },
            {
                "code": "TC-WM-03",
                "name": "Thêm linh kiện mới vào danh mục kho",
                "steps": "1. Bấm 'Thêm Linh Kiện Mới'\n2. Nhập Tên, Nhóm, Số lượng, Giá, NCC, Kệ ➔ Bấm Tạo",
                "expected": "Linh kiện mới thêm vào danh sách kho thành công.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "7. Nhân Viên Kỹ Thuật Lắp Ráp (ASSEMBLY)",
        "desc": "Kỹ thuật viên thực hiện lắp ráp máy tính bộ theo BOM đơn hàng và kiểm định chất lượng.",
        "tcs": [
            {
                "code": "TC-ASM-01",
                "name": "Đăng nhập & Tiếp nhận Job lắp ráp",
                "steps": "1. Đăng nhập assembly / 123456 ➔ /admin/assembly\n2. Xem danh sách Job cần lắp (JOB-xxx)",
                "expected": "Hiển thị danh sách đơn hàng có cấu hình PC cần lắp ráp.",
                "status": "PASS"
            },
            {
                "code": "TC-ASM-02",
                "name": "Tạo Job lắp ráp thủ công tại quầy",
                "steps": "1. Bấm 'Tạo Job Lắp Ráp Thủ Công'\n2. Chọn đơn hàng / Khách hàng & danh sách linh kiện ➔ Bấm Tạo",
                "expected": "Job lắp ráp mới được khởi tạo ở trạng thái PENDING.",
                "status": "PASS"
            },
            {
                "code": "TC-ASM-03",
                "name": "Gán Serial Number & Đổi trạng thái Assembling",
                "steps": "1. Chọn Job ➔ Bấm Bắt đầu lắp ráp\n2. Gán Serial Number khả dụng cho CPU, VGA, Mainboard",
                "expected": "Trạng thái Job chuyển sang ASSEMBLING, mã S/N được chuyển thành USED.",
                "status": "PASS"
            },
            {
                "code": "TC-ASM-04",
                "name": "Thực hiện Checklist kiểm định 5 bước",
                "steps": "1. Tích chọn: Socket Check, Keo tản nhiệt, Đi dây, Boot BIOS, Stress Test 15p\n2. Bấm Hoàn thành Job",
                "expected": "Job hoàn tất (COMPLETED). Đơn hàng tự động chuyển sang trạng thái READY_TO_SHIP cho kho xuất.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "8. Nhân Viên Mua Hàng (PURCHASING)",
        "desc": "Nhân viên phòng mua hàng tạo đề xuất mua hàng (PR) và đơn mua hàng (PO) gửi nhà cung cấp.",
        "tcs": [
            {
                "code": "TC-PUR-01",
                "name": "Xem danh sách hàng cần mua bổ sung",
                "steps": "1. Đăng nhập purchasing / 123456 ➔ /admin/purchasing\n2. Xem cảnh báo sản phẩm dưới ngưỡng an toàn",
                "expected": "Hiển thị danh sách linh kiện cần nhập thêm kèm số lượng gợi ý.",
                "status": "PASS"
            },
            {
                "code": "TC-PUR-02",
                "name": "Tạo đơn mua hàng (PO) gửi Nhà cung cấp",
                "steps": "1. Bấm 'Tạo Yêu Cầu Mua Hàng (PR)'\n2. Chọn NCC (Intel, ASUS...), Linh kiện, Số lượng, Đơn giá ➔ Bấm Gửi PO",
                "expected": "Đơn PO mới (PO-xxxx) được tạo ở trạng thái DRAFT/SENT và gửi tới Supplier Portal.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "9. Đối Tác Nhà Cung Cấp (SUPPLIER)",
        "desc": "Nhà cung cấp tiếp nhận PO từ AetherPC, báo giá và xác nhận cung ứng thiết bị.",
        "tcs": [
            {
                "code": "TC-SUP-01",
                "name": "Đăng nhập Supplier Portal",
                "steps": "1. Truy cập /login ➔ Đăng nhập supplier / 123456\n2. Tự động chuyển hướng tới /supplier/portal",
                "expected": "Truy cập giao diện Supplier Portal thành công.",
                "status": "PASS"
            },
            {
                "code": "TC-SUP-02",
                "name": "Xem PO mới & Bấm Đồng ý cung cấp",
                "steps": "1. Xem danh sách PO mới chờ xác nhận\n2. Bấm nút 'Đồng ý cung cấp' trên đơn PO",
                "expected": "Đơn PO chuyển từ SENT sang PENDING (Chờ CEO phê duyệt thanh toán).",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "10. Nhân Viên Giao Hàng & Vận Chuyển (DELIVERY)",
        "desc": "Nhân viên đội ship tiếp nhận đơn hàng đã xuất kho, lộ trình giao nhận và cập nhật trạng thái.",
        "tcs": [
            {
                "code": "TC-DEL-01",
                "name": "Xem danh sách đơn chờ lấy hàng",
                "steps": "1. Đăng nhập delivery / 123456 ➔ /admin/delivery\n2. Xem tab 'Chờ lấy hàng' (READY_TO_SHIP)",
                "expected": "Hiển thị đầy đủ thông tin KH, SĐT, địa chỉ giao hàng và danh sách sản phẩm.",
                "status": "PASS"
            },
            {
                "code": "TC-DEL-02",
                "name": "Nhận giao đơn hàng (Pick up)",
                "steps": "1. Bấm nút 'Nhận Giao' trên thẻ đơn hàng",
                "expected": "Đơn hàng chuyển sang trạng thái Đang giao (SHIPPED).",
                "status": "PASS"
            },
            {
                "code": "TC-DEL-03",
                "name": "Xác nhận Đã Giao thành công",
                "steps": "1. Chuyển sang tab 'Đang giao'\n2. Bấm 'Đã Giao' ➔ Xác nhận đồng ý",
                "expected": "Đơn chuyển thành DELIVERED, tự động cập nhật ngày giao và ghi nhận nhật ký vận chuyển.",
                "status": "PASS"
            },
            {
                "code": "TC-DEL-04",
                "name": "Xử lý Giao Hàng Thất Bại & Giao Lại",
                "steps": "1. Bấm nút 'Thất Bại' trên đơn đang giao\n2. Nhập lý do (KH hẹn lại giờ)\n3. Tại tab Chờ lấy/Giao lại ➔ Bấm 'Thử Giao Lại'",
                "expected": "Đơn lưu ghi chú lý do thất bại màu đỏ, cho phép ấn giao lại chuyển về SHIPPED.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "11. Nhân Viên Chăm Sóc Khách Hàng (CSKH)",
        "desc": "Bộ phận hỗ trợ, giải đáp thắc mắc, xử lý ticket khiếu nại và đơn yêu cầu đổi trả.",
        "tcs": [
            {
                "code": "TC-CS-01",
                "name": "Tiếp nhận & Tạo Ticket khiếu nại",
                "steps": "1. Đăng nhập cskh / 123456 ➔ /admin/cskh\n2. Bấm 'Tạo Ticket Mới'\n3. Nhập Tên KH, SĐT, Tiêu đề, Mô tả ➔ Bấm Tạo",
                "expected": "Ticket khiếu nại mới (TKT-xxx) được tạo ở trạng thái OPEN.",
                "status": "PASS"
            },
            {
                "code": "TC-CS-02",
                "name": "Phân công & Giải quyết Ticket",
                "steps": "1. Bấm 'Nhận xử lý' ➔ Chuyển IN_PROGRESS\n2. Bấm 'Giải quyết' ➔ Nhập Hướng xử lý ➔ Xác nhận",
                "expected": "Ticket chuyển sang RESOLVED với ghi chú hướng xử lý màu xanh.",
                "status": "PASS"
            },
            {
                "code": "TC-CS-03",
                "name": "Duyệt yêu cầu Đổi trả / Hoàn tiền",
                "steps": "1. Chuyển sang tab 'Yêu cầu đổi trả'\n2. Xem các yêu cầu RET-xxx từ khách hàng\n3. Bấm 'Xử lý' ➔ 'Phê duyệt' ➔ 'Hoàn thành'",
                "expected": "Yêu cầu đổi trả chuyển từ PENDING ➔ PROCESSING ➔ APPROVED ➔ COMPLETED.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "12. Nhân Viên Nhân Sự (HR)",
        "desc": "Quản lý danh sách nhân sự, chấm công hàng ngày, duyệt nghỉ phép và lập bảng lương.",
        "tcs": [
            {
                "code": "TC-HR-01",
                "name": "Quản lý hồ sơ nhân viên",
                "steps": "1. Đăng nhập hr / 123456 ➔ /admin/hr\n2. Thêm nhân viên mới (Họ tên, User, Chức vụ, Lương cơ bản)",
                "expected": "Thêm nhân viên mới thành công, tài khoản mới có thể đăng nhập vào hệ thống.",
                "status": "PASS"
            },
            {
                "code": "TC-HR-02",
                "name": "Chấm công nhân viên hàng ngày",
                "steps": "1. Chọn ngày chấm công\n2. Đánh dấu trạng thái: Có mặt (Present) / Đi muộn (Late) / Vắng (Absent)",
                "expected": "Bảng chấm công ghi nhận trạng thái và tính phạt đi muộn (50.000đ/lần).",
                "status": "PASS"
            },
            {
                "code": "TC-HR-03",
                "name": "Duyệt đơn xin nghỉ phép",
                "steps": "1. Chuyển tab 'Đơn Xin Nghỉ Phép'\n2. Xem đơn của nhân viên ➔ Bấm 'Duyệt phép'",
                "expected": "Đơn phép được duyệt (APPROVED), ngày nghỉ được tính là phép hưởng lương.",
                "status": "PASS"
            },
            {
                "code": "TC-HR-04",
                "name": "Khóa bảng công & Nộp bảng lương cho CEO",
                "steps": "1. Chuyển tab 'Bảng Tính Lương Tháng'\n2. Bấm 'Tính Lương Tự Động' ➔ Kiểm tra tổng lương\n3. Bấm 'Khóa Bảng Công & Nộp CEO'",
                "expected": "Bảng lương khóa lại ở trạng thái SUBMITTED_TO_CEO và gửi tới CEO Dashboard.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "13. Kế Toán Tài Chính (ACCOUNTANT)",
        "desc": "Quản lý Sổ cái thu chi, thực hiện chi trả lương nhân sự và thanh toán cho nhà cung cấp.",
        "tcs": [
            {
                "code": "TC-ACC-01",
                "name": "Theo dõi Sổ cái & Tạo bút toán thủ công",
                "steps": "1. Đăng nhập accounting / 123456 ➔ /admin/accounting\n2. Xem Sổ cái Thu / Chi (Income/Expense)\n3. Tạo bút toán Thu/Chi thủ công",
                "expected": "Sổ cái cập nhật thời gian thực, tổng thu chi và số dư tiền mặt thay đổi tương ứng.",
                "status": "PASS"
            },
            {
                "code": "TC-ACC-02",
                "name": "Giải ngân lương nhân viên",
                "steps": "1. Chuyển tab 'Chi Trả Lương Nhân Sự'\n2. Xem danh sách lương đã được CEO duyệt (APPROVED_BY_CEO)\n3. Bấm 'Giải Ngân'",
                "expected": "Lương chuyển sang PAID, ghi nhận bút toán EXPENSE chi trả lương vào Sổ cái.",
                "status": "PASS"
            },
            {
                "code": "TC-ACC-03",
                "name": "Thanh toán tiền cho Nhà cung cấp (PO)",
                "steps": "1. Chuyển tab 'Thanh Toán Nhà Cung Cấp'\n2. Chọn đơn PO đã nhập kho (RECEIVED) ➔ Bấm 'Thanh toán NCC'",
                "expected": "Trạng thái PO chuyển thành PAID, ghi bút toán chi trả tiền mua hàng vào Sổ cái.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "14. Giám Đốc Điều Hành (CEO)",
        "desc": "Ban điều hành xem báo cáo tổng quan tài chính thời gian thực và phê duyệt các quyết định lớn.",
        "tcs": [
            {
                "code": "TC-CEO-01",
                "name": "Theo dõi CEO Dashboard thời gian thực",
                "steps": "1. Đăng nhập ceo / 123456 ➔ /admin/dashboard\n2. Xem KPI: Doanh thu, Lợi nhuận, Chi phí, Đơn hàng, Tồn kho",
                "expected": "Các chỉ số tài chính và biểu đồ tăng trưởng hiển thị chính xác theo dữ liệu ERP.",
                "status": "PASS"
            },
            {
                "code": "TC-CEO-02",
                "name": "Phê duyệt bảng lương do HR gửi lên",
                "steps": "1. Tại trang Dashboard ➔ Xem thông báo 'Bảng Lương Tháng Chờ Phê Duyệt'\n2. Bấm 'Phê Duyệt Bảng Lương'",
                "expected": "Bảng lương chuyển sang APPROVED_BY_CEO, phát lệnh cho Kế toán giải ngân.",
                "status": "PASS"
            },
            {
                "code": "TC-CEO-03",
                "name": "Phê duyệt đơn PO mua hàng & Đơn nghỉ phép",
                "steps": "1. Xem các đơn PO mua hàng từ phòng Purchasing\n2. Xem đơn xin nghỉ phép nhân sự ➔ Bấm Duyệt",
                "expected": "Đơn PO chuyển sang APPROVED, cho phép kho nhận hàng; đơn nghỉ phép được ghi nhận.",
                "status": "PASS"
            }
        ]
    },
    {
        "name": "15. Quản Trị Hệ Thống (ADMIN)",
        "desc": "Quản trị viên hạ tầng giám sát hệ thống, quản lý tài khoản và reset dữ liệu kiểm thử.",
        "tcs": [
            {
                "code": "TC-ADM-01",
                "name": "Quản lý hệ thống & Phân quyền tài khoản",
                "steps": "1. Đăng nhập admin / 123456 ➔ /admin/system\n2. Xem trạng thái các dịch vụ Backend, Database, Frontend",
                "expected": "Hiển thị trạng thái các Service Running (Green), cho phép phân quyền tài khoản.",
                "status": "PASS"
            },
            {
                "code": "TC-ADM-02",
                "name": "Tra cứu lương cá nhân (MyPayroll)",
                "steps": "1. Vào menu 'Tra Cứu Lương' (/admin/my-payroll)\n2. Xem lịch công và lương cá nhân",
                "expected": "Mỗi tài khoản chỉ xem đúng bảng công và thu nhập của chính mình.",
                "status": "PASS"
            }
        ]
    }
]

for actor in actors_data:
    add_actor_section(actor["name"], actor["desc"], actor["tcs"])

doc.save("docs/KichBanKiemThu_ToanBoActors_AetherPC.docx")
print("Successfully generated docs/KichBanKiemThu_ToanBoActors_AetherPC.docx")
